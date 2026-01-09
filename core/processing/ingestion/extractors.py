"""Helper functions for extracting document contents in ingestion pipeline."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Union

from ..common import logger
from ..errors import ProcessingPipelineError


def extract_enhanced_docx_content(file_path: Union[str, Path]) -> str:
    """Extract DOCX content with tables, metadata, and headings preserved."""
    try:
        import docx
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = docx.Document(file_path)
        content_parts: list[str] = []

        for element in doc.element.body:
            if element.tag.endswith("p"):
                para = Paragraph(element, doc)
                if para.text.strip():
                    style_name = para.style.name if para.style else "Normal"
                    if "Heading" in style_name:
                        content_parts.append(f"\n## {para.text.strip()}\n")
                    else:
                        content_parts.append(para.text.strip())
            elif element.tag.endswith("tbl"):
                table = Table(element, doc)
                table_content = _extract_table_content(table)
                if table_content:
                    content_parts.append(f"\n[ТАБЛИЦА]\n{table_content}\n[/ТАБЛИЦА]\n")

        metadata_text = _extract_docx_metadata(doc)
        if metadata_text:
            content_parts.insert(0, f"[МЕТАДАННЫЕ]\n{metadata_text}\n[/МЕТАДАННЫЕ]\n")

        return "\n".join(content_parts)

    except Exception as exc:  # pragma: no cover - fallbacks rely on optional deps
        logger.warning(
            "Ошибка улучшенного извлечения DOCX %s: %s, используем простое извлечение",
            file_path,
            exc,
        )
        try:
            import docx

            doc = docx.Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as fallback_error:  # pragma: no cover - defensive branch
            raise ProcessingPipelineError(
                f"DOCX extraction failed: {fallback_error}"
            ) from fallback_error


def extract_doc_content(file_path: Union[str, Path]) -> str:
    """Extract legacy DOC content using best available tools."""
    try:
        try:
            import docx2txt

            content = docx2txt.process(str(file_path))
            if content and content.strip():
                return content.strip()
        except ImportError:
            logger.warning("docx2txt не установлен, пробуем альтернативные методы")

        try:
            import textract

            content = textract.process(str(file_path)).decode("utf-8")
            if content and content.strip():
                return content.strip()
        except ImportError:
            logger.warning("textract не установлен")
        except Exception as exc:  # pragma: no cover - textract optional
            logger.warning("Ошибка textract для DOC: %s", exc)

        try:
            import olefile

            if olefile.isOleFile(str(file_path)):
                logger.info("Обнаружен OLE2 файл, используем базовое извлечение")
                return (
                    f"[DOC ФАЙЛ: {Path(file_path).name}]\n"
                    "Содержимое legacy DOC файла требует специальных инструментов для извлечения текста."
                )
        except ImportError:
            pass

        raise ProcessingPipelineError(
            "Не удалось найти подходящий инструмент для обработки DOC файла"
        )

    except Exception as exc:
        logger.error("Ошибка извлечения DOC содержимого: %s", exc)
        raise ProcessingPipelineError(f"DOC extraction failed: {exc}") from exc


def extract_rtf_content(file_path: Union[str, Path]) -> str:
    """Extract text from RTF documents."""
    try:
        try:
            from striprtf.striprtf import rtf_to_text

            with open(file_path, "r", encoding="utf-8", errors="ignore") as source:
                rtf_content = source.read()

            text_content = rtf_to_text(rtf_content)
            if text_content and text_content.strip():
                return text_content.strip()
        except ImportError:
            logger.warning("striprtf не установлен, пробуем альтернативные методы")
        except Exception as exc:  # pragma: no cover - optional dependency
            logger.warning("Ошибка striprtf: %s", exc)

        try:
            import textract

            content = textract.process(str(file_path)).decode("utf-8")
            if content and content.strip():
                return content.strip()
        except ImportError:
            logger.warning("textract не установлен")
        except Exception as exc:  # pragma: no cover - optional dependency
            logger.warning("Ошибка textract для RTF: %s", exc)

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as source:
                raw_content = source.read()

            import re

            clean_content = re.sub(r"\{[^}]*\}", "", raw_content)
            clean_content = re.sub(r"\\[a-z]+\d*\s?", "", clean_content)
            clean_content = re.sub(r"\s+", " ", clean_content).strip()

            if clean_content:
                return clean_content
        except Exception as fallback_error:
            logger.warning("Ошибка базового RTF извлечения: %s", fallback_error)

        raise ProcessingPipelineError("Не удалось извлечь текст из RTF файла")

    except Exception as exc:
        logger.error("Ошибка извлечения RTF содержимого: %s", exc)
        raise ProcessingPipelineError(f"RTF extraction failed: {exc}") from exc


def extract_spreadsheet_content(file_path: Union[str, Path], extension: str) -> str:
    """Extract structured content from spreadsheet formats (Excel/CSV)."""
    try:
        import pandas as pd

        resolved_path = Path(file_path)
        content_parts: list[str] = []

        if extension == ".csv":
            _extract_csv_content(resolved_path, content_parts, pd)
        elif extension in [".xlsx", ".xls"]:
            _extract_excel_content(resolved_path, extension, content_parts, pd)

        content_parts.append("")
        content_parts.append("[СТАТИСТИКА]")
        content_parts.append(f"Файл: {resolved_path.name}")
        content_parts.append(f"Формат: {extension.upper()}")
        content_parts.append(
            f"Извлечено текстовых блоков: {len([p for p in content_parts if p.strip()])}"
        )
        content_parts.append(
            f"Обработано: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )

        result = "\n".join(content_parts)
        if not result.strip():
            raise ProcessingPipelineError(f"No content extracted from {extension} file")

        logger.info(
            "✅ Извлечено %s символов из %s файла %s",
            len(result),
            extension,
            resolved_path.name,
        )
        return result

    except ImportError:
        logger.warning("pandas не доступен, используем базовый метод извлечения")
        return _extract_spreadsheet_without_pandas(file_path, extension)
    except Exception as exc:
        logger.error("❌ Ошибка извлечения из %s файла: %s", extension, exc)
        raise ProcessingPipelineError(
            f"Spreadsheet extraction failed: {exc}"
        ) from exc


def _extract_csv_content(path: Path, content_parts: list[str], pd) -> None:
    encodings = ["utf-8", "cp1251", "latin-1", "iso-8859-1"]
    dataframe = None

    for encoding in encodings:
        try:
            dataframe = pd.read_csv(path, encoding=encoding)
            logger.info("✅ CSV файл успешно прочитан с кодировкой %s", encoding)
            break
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            if encoding == encodings[-1]:
                raise exc

    if dataframe is None:
        raise ProcessingPipelineError("Не удалось прочитать CSV файл с доступными кодировками")

    content_parts.append(f"[CSV ФАЙЛ: {path.name}]")
    content_parts.append(f"Размер таблицы: {len(dataframe)} строк, {len(dataframe.columns)} столбцов")

    if not dataframe.columns.empty:
        content_parts.append("[ЗАГОЛОВКИ СТОЛБЦОВ]")
        content_parts.append(" | ".join(str(col) for col in dataframe.columns))
        content_parts.append("")

    max_rows = min(len(dataframe), 100)
    content_parts.append(f"[ДАННЫЕ ТАБЛИЦЫ - первые {max_rows} строк]")

    for idx, row in dataframe.head(max_rows).iterrows():
        row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row.values)
        if row_text.strip():
            content_parts.append(f"Строка {idx + 1}: {row_text}")

    if len(dataframe) > 100:
        content_parts.append(f"... и еще {len(dataframe) - 100} строк")


def _extract_excel_content(path: Path, extension: str, content_parts: list[str], pd) -> None:
    try:
        excel_data = pd.read_excel(
            path,
            sheet_name=None,
            engine="openpyxl" if extension == ".xlsx" else "xlrd",
        )

        content_parts.append(f"[EXCEL ФАЙЛ: {path.name}]")
        content_parts.append(f"Количество листов: {len(excel_data)}")
        content_parts.append("")

        for sheet_name, dataframe in excel_data.items():
            if dataframe.empty:
                continue

            content_parts.append(f"[ЛИСТ: {sheet_name}]")
            content_parts.append(
                f"Размер: {len(dataframe)} строк, {len(dataframe.columns)} столбцов"
            )

            if not dataframe.columns.empty:
                content_parts.append(
                    "Столбцы: " + " | ".join(str(col) for col in dataframe.columns)
                )
                content_parts.append("")

            max_rows = min(len(dataframe), 50)
            for row in dataframe.head(max_rows).itertuples(index=False):
                row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row)
                if row_text.strip():
                    content_parts.append(row_text)

            if len(dataframe) > 50:
                content_parts.append(f"... и еще {len(dataframe) - 50} строк")

            content_parts.append("")

    except Exception as exc:
        logger.warning("⚠️ Ошибка pandas для Excel: %s", exc)
        if extension == ".xlsx":
            _extract_excel_with_openpyxl(path, content_parts)
        else:
            raise ProcessingPipelineError(f"XLS processing failed: {exc}") from exc


def _extract_excel_with_openpyxl(path: Path, content_parts: list[str]) -> None:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(path, data_only=True)
        content_parts.append(f"[EXCEL ФАЙЛ (openpyxl): {path.name}]")

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            content_parts.append(f"[ЛИСТ: {sheet_name}]")

            for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                if row_num > 100:
                    break

                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    content_parts.append(row_text)

            content_parts.append("")

    except Exception as exc:  # pragma: no cover - fallbacks optional
        raise ProcessingPipelineError(f"Excel processing failed: {exc}") from exc


def _extract_spreadsheet_without_pandas(file_path: Union[str, Path], extension: str) -> str:
    if extension == ".xlsx":
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(file_path)
            text_parts: list[str] = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"[ЛИСТ: {sheet_name}]")

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except ImportError as exc:
            raise ProcessingPipelineError("openpyxl не установлен для обработки Excel файлов") from exc

    if extension == ".csv":
        try:
            import csv

            text_parts: list[str] = []

            with open(file_path, "r", encoding="utf-8", newline="") as csvfile:
                reader = csv.reader(csvfile)
                for row_num, row in enumerate(reader, 1):
                    if row_num > 100:
                        break
                    row_text = " | ".join(row)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except Exception as exc:  # pragma: no cover - defensive fallback
            raise ProcessingPipelineError(f"CSV processing failed: {exc}") from exc

    raise ProcessingPipelineError(f"No fallback available for {extension} without pandas")


def _extract_table_content(table) -> str:
    try:
        lines = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as exc:
        logger.debug("Ошибка извлечения таблицы: %s", exc)
        return ""


def _extract_docx_metadata(doc) -> str:
    try:
        core_props = doc.core_properties
        metadata_parts = []

        if core_props.title:
            metadata_parts.append(f"Название: {core_props.title}")
        if core_props.author:
            metadata_parts.append(f"Автор: {core_props.author}")
        if core_props.created:
            metadata_parts.append(
                f"Дата создания: {core_props.created.strftime('%Y-%m-%d %H:%M:%S')}"
            )
        if core_props.modified:
            metadata_parts.append(
                f"Дата изменения: {core_props.modified.strftime('%Y-%m-%d %H:%M:%S')}"
            )
        if core_props.last_modified_by:
            metadata_parts.append(
                f"Последний редактор: {core_props.last_modified_by}"
            )
        if core_props.subject:
            metadata_parts.append(f"Тема: {core_props.subject}")
        if core_props.category:
            metadata_parts.append(f"Категория: {core_props.category}")

        return "\n".join(metadata_parts)
    except Exception as exc:
        logger.debug("Ошибка извлечения метаданных DOCX: %s", exc)
        return ""
